"""End-to-end extraction test against a synthetic .docx fixture.

Skipped automatically if python-docx isn't installed in the environment.
"""

import json
import tempfile
import unittest
from pathlib import Path

try:
    import docx  # type: ignore
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class ExtractionEndToEndTests(unittest.TestCase):

    def _make_fixture(self) -> str:
        from docx import Document
        doc = Document()
        # Real manuscripts style chapter headings as Heading 1, which is what
        # our chapter detector keys off of.
        doc.add_paragraph("Chapter 1: The Names We Carry (ACT I)", style="Heading 1")
        doc.add_paragraph(
            "Ramos walked the line at Khe Sanh. The 1st Cavalry was dug in. "
            "He whispered, “Cuídate, mijo.”"
        )
        doc.add_paragraph("* * *")  # space-separated scene break, as in Ramos
        doc.add_paragraph(
            "Mi abuela said the rosary, the way she always did when soldiers left. "
            "Órale, the LT said. Move out."
        )
        doc.add_paragraph("Chapter 2", style="Heading 1")
        doc.add_paragraph(
            "Two years later, the casualty list arrived. Twelve names. Half of them Hispanic."
        )

        path = Path(tempfile.mkdtemp()) / "fixture.docx"
        doc.save(str(path))
        return str(path)

    def test_extract_preserves_text_and_assigns_chapters(self):
        from elasticsearch.pipelines.extract_manuscript import extract_manuscript

        fixture = self._make_fixture()
        out_path = Path(fixture).with_suffix(".json")
        docs = extract_manuscript(fixture, str(out_path))

        self.assertGreaterEqual(len(docs), 3)
        chapter_seqs = sorted({d["chapter_seq"] for d in docs})
        self.assertEqual(chapter_seqs, [1, 2])
        labels = sorted({d["chapter_label_num"] for d in docs if d["chapter_label_num"] is not None})
        self.assertEqual(labels, [1, 2])

        # Act parsed from the heading text propagates to Ch1 paragraphs.
        act_i_paras = [d for d in docs if d["chapter_seq"] == 1]
        self.assertTrue(all(d["act"] == "Act I" for d in act_i_paras))

        # The em-dash and Spanish accent must be intact.
        all_text = "\n".join(d["text"] for d in docs)
        self.assertIn("Cuídate, mijo.", all_text)
        self.assertIn("Órale", all_text)

        # Space-separated scene break must NOT appear as a paragraph doc.
        self.assertFalse(any(d["text"].strip() == "* * *" for d in docs))

        # Every paragraph has the required forensic fields.
        for d in docs:
            self.assertIn("text_hash", d)
            self.assertIn("doc_id", d)
            self.assertIn("pipeline_version", d)
            self.assertIn("extraction_timestamp", d)
            self.assertIn("chapter_seq", d)
            self.assertIn("chapter_label_num", d)
            self.assertEqual(d["doc_type"], "paragraph")

        # The on-disk JSON round-trips.
        with open(out_path, "r", encoding="utf-8") as f:
            loaded = json.load(f)
        self.assertEqual(len(loaded), len(docs))


if __name__ == "__main__":
    unittest.main()
